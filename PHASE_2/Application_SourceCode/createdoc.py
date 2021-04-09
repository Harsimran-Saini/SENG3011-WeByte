from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Correlation report', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_page_break()

document.add_heading('Correlation scatter plot', level=1)
document.add_paragraph('Insert plot here', style='Intense Quote')
document.add_picture('sample.png', width=Inches(1.25))

document.add_heading('Correlation map', level=1)
document.add_paragraph('Insert map here', style='Intense Quote')
document.add_picture('sample.png', width=Inches(1.25))

document.add_page_break()

document.add_heading('Bibliography', level=1)
document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

document.add_page_break()
document.add_heading('Appendix', level=1)
records = (
    ('Australia', '50', '41'),
    ('US', '500', '415'),
    ('Canada', '54', '400')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Country'
hdr_cells[1].text = 'Covid Cases'
hdr_cells[2].text = 'Covid News Articles'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc


document.save('demo.docx')